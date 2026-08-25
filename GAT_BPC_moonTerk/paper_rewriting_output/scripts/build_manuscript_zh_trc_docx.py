#!/usr/bin/env python3
"""Build a publication-review DOCX from manuscript_zh_trc.md.

The exporter keeps equations as native Word OMML objects.  Display equations
are placed in borderless three-column tables so the equation remains centred
while its explicit manuscript number is right aligned.
"""

from __future__ import annotations

import argparse
import re
import shutil
import tempfile
import zipfile
from datetime import date
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from lxml import etree


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_ROOT = PROJECT_ROOT / "paper_rewriting_output"
DEFAULT_SOURCE = OUTPUT_ROOT / "manuscript_zh_trc.md"
DEFAULT_OUTPUT = OUTPUT_ROOT / "final_paper" / "manuscript_zh_trc.docx"
DEFAULT_REPORT = OUTPUT_ROOT / "word_report.md"

TITLE_TEXT = "面向月球水冰探测多趟次路线规划的学习引导精确分支定价切割算法"
EQUATION_MARKER_PREFIX = "PS_EQNO_"
EQUATION_MARKER_SUFFIX = "_END"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": W_NS, "m": M_NS}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--report", type=Path, default=DEFAULT_REPORT)
    return parser.parse_args()


def set_style_font(style, *, size_pt: float, east_asia: str, latin: str = "Times New Roman", bold: bool | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_run_font(run, *, size_pt: float | None = None, east_asia: str = "宋体", latin: str = "Times New Roman") -> None:
    run.font.name = latin
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def ensure_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("第 ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    paragraph.add_run(" 页")
    for footer_run in paragraph.runs:
        set_run_font(footer_run, size_pt=9.0)


def configure_page(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def build_reference_docx(path: Path) -> None:
    document = Document()
    configure_page(document.sections[0])

    normal = document.styles["Normal"]
    set_style_font(normal, size_pt=12.0, east_asia="宋体")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    body = ensure_style(document, "Body Text")
    body.base_style = normal
    set_style_font(body, size_pt=12.0, east_asia="宋体")
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.first_line_indent = Cm(0.74)
    body.paragraph_format.space_after = Pt(0)

    first = ensure_style(document, "First Paragraph")
    first.base_style = body
    set_style_font(first, size_pt=12.0, east_asia="宋体")
    first.paragraph_format.first_line_indent = Cm(0.74)

    compact = ensure_style(document, "Compact")
    compact.base_style = body
    set_style_font(compact, size_pt=11.5, east_asia="宋体")
    compact.paragraph_format.left_indent = Cm(0.74)
    compact.paragraph_format.first_line_indent = Cm(-0.42)
    compact.paragraph_format.line_spacing = 1.25
    compact.paragraph_format.space_after = Pt(4)

    title = document.styles["Title"]
    set_style_font(title, size_pt=18.0, east_asia="黑体", bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.keep_with_next = True

    heading_specs = (
        ("Heading 1", 14.0, 18, 8),
        ("Heading 2", 12.5, 14, 6),
        ("Heading 3", 12.0, 10, 4),
        ("Heading 4", 11.5, 8, 4),
    )
    for name, size, before, after in heading_specs:
        style = document.styles[name]
        set_style_font(style, size_pt=size, east_asia="黑体", bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = ensure_style(document, "Caption")
    set_style_font(caption, size_pt=9.5, east_asia="宋体")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    for style_name in ("Image Caption", "Table Caption"):
        style = ensure_style(document, style_name)
        style.base_style = caption
        set_style_font(style, size_pt=9.5, east_asia="宋体")
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    footer_p = footer.paragraphs[0]
    footer_p.clear()
    add_page_number(footer_p)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


DISPLAY_EQUATION_RE = re.compile(
    r"\$\$\s*(?P<body>.*?)\s*\\tag\{(?P<tag>[^{}]+)\}\s*\$\$",
    flags=re.DOTALL,
)


def normalize_working_citations(text: str) -> tuple[str, int]:
    pattern = re.compile(r"\[@(?P<first>C\d+)(?P<rest>(?:\s*;\s*@C\d+)*)\]")

    def replace(match: re.Match[str]) -> str:
        keys = [match.group("first")]
        keys.extend(re.findall(r"C\d+", match.group("rest")))
        return "[" + "; ".join(keys) + "]"

    normalized, count = pattern.subn(replace, text)
    normalized = normalized.replace("[@Cxxx]", "[Cxxx]")
    return normalized, count


def preprocess_markdown(source: Path, destination: Path) -> dict[str, int]:
    text = source.read_text(encoding="utf-8")
    citation_text, citation_count = normalize_working_citations(text)
    equation_count = 0

    def replace_equation(match: re.Match[str]) -> str:
        nonlocal equation_count
        equation_count += 1
        body = match.group("body").strip()
        tag = match.group("tag").strip()
        marker = f"{EQUATION_MARKER_PREFIX}{equation_count:04d}_{tag}{EQUATION_MARKER_SUFFIX}"
        return f"$$\n{body}\n$$\n\n{marker}"

    processed = DISPLAY_EQUATION_RE.sub(replace_equation, citation_text)
    destination.write_text(processed, encoding="utf-8")
    return {
        "source_chars": len(text),
        "equation_count": equation_count,
        "working_citation_groups": citation_count,
        "markdown_table_count": sum(1 for line in text.splitlines() if line.startswith("|---")),
        "image_count": text.count("!["),
        "heading_count": sum(1 for line in text.splitlines() if line.startswith("#")),
    }


def set_cell_margins(cell, *, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def table_widths(column_count: int) -> list[float]:
    if column_count == 2:
        return [0.75, 5.45]
    if column_count == 3:
        return [1.10, 2.05, 3.05]
    if column_count == 6:
        return [1.10, 0.75, 0.82, 0.95, 1.20, 1.38]
    if column_count == 7:
        return [1.08, 0.68, 0.92, 0.92, 0.88, 0.88, 0.84]
    return [6.20 / max(column_count, 1)] * column_count


def style_tables(document: Document) -> None:
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_fixed_layout(table)
        if table.rows:
            set_repeat_table_header(table.rows[0])
        widths = table_widths(len(table.columns))
        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            for column_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if column_index < len(widths):
                    cell.width = Inches(widths[column_index])
                if row_index == 0:
                    shade_cell(cell, "E8EDF0")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.left_indent = Cm(0)
                    paragraph.paragraph_format.right_indent = Cm(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.keep_together = True
                    if row_index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif len(table.columns) >= 6 or column_index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_run_font(run, size_pt=9.0)
                        if row_index == 0:
                            run.bold = True


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def style_document(path: Path) -> None:
    document = Document(path)
    for section in document.sections:
        configure_page(section)
        section.start_type = WD_SECTION.NEW_PAGE if section is not document.sections[0] else section.start_type
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_p = footer.paragraphs[0]
        footer_p.clear()
        add_page_number(footer_p)

    document.core_properties.title = TITLE_TEXT
    document.core_properties.subject = "TRC 风格中文审阅稿"
    document.core_properties.keywords = "月球水冰探测；多路径；多趟次；分支定价切割；SPPRC"
    document.core_properties.comments = "由 manuscript_zh_trc.md 自动生成；公式为 Word 原生 OMML。"

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        paragraph.paragraph_format.widow_control = True

        if text == TITLE_TEXT:
            paragraph.style = document.styles["Title"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, size_pt=18.0, east_asia="黑体")
                run.bold = True
            continue

        if style_name.startswith("Heading"):
            level_match = re.search(r"(\d+)$", style_name)
            level = int(level_match.group(1)) if level_match else 1
            east_font = "黑体"
            size = {1: 14.0, 2: 12.5, 3: 12.0}.get(level, 11.5)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            if text.startswith(("附录 A", "参考文献说明")):
                paragraph.paragraph_format.page_break_before = True
            for run in paragraph.runs:
                set_run_font(run, size_pt=size, east_asia=east_font)
                run.bold = True
            continue

        is_caption = bool(re.match(r"^(?:表|图)\s*\d+\u3000", text)) or "Caption" in style_name
        if is_caption:
            paragraph.style = document.styles["Caption"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.keep_together = True
            for run in paragraph.runs:
                set_run_font(run, size_pt=9.5)
            continue

        if paragraph_has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_together = True
            continue

        if not text:
            continue

        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        if style_name == "Compact":
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.first_line_indent = Cm(-0.42)
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = Pt(4)
        if text.startswith(("关键词：", "输入：", "输出：", "注：")):
            paragraph.paragraph_format.first_line_indent = Cm(0)
        if text.startswith("关键词："):
            paragraph.paragraph_format.space_after = Pt(10)
        if text.startswith("注："):
            paragraph.paragraph_format.line_spacing = 1.2
            for run in paragraph.runs:
                set_run_font(run, size_pt=10.0)
            continue
        for run in paragraph.runs:
            set_run_font(run, size_pt=12.0)

    style_tables(document)

    max_width = Inches(6.20)
    for shape in document.inline_shapes:
        if shape.width > max_width:
            ratio = shape.height / shape.width
            shape.width = max_width
            shape.height = int(max_width * ratio)

    document.save(path)


def w_element(tag: str, **attributes: str):
    element = etree.Element(f"{{{W_NS}}}{tag}")
    for key, value in attributes.items():
        element.set(f"{{{W_NS}}}{key}", value)
    return element


def make_cell(width: int, *, paragraph=None, number_text: str | None = None):
    tc = w_element("tc")
    tc_pr = w_element("tcPr")
    tc_pr.append(w_element("tcW", w=str(width), type="dxa"))
    tc_pr.append(w_element("vAlign", val="center"))
    margins = w_element("tcMar")
    for side in ("top", "start", "bottom", "end"):
        margins.append(w_element(side, w="30", type="dxa"))
    tc_pr.append(margins)
    tc.append(tc_pr)

    if paragraph is not None:
        p_pr = paragraph.find(f"{{{W_NS}}}pPr")
        if p_pr is None:
            p_pr = w_element("pPr")
            paragraph.insert(0, p_pr)
        for old_jc in p_pr.findall(f"{{{W_NS}}}jc"):
            p_pr.remove(old_jc)
        p_pr.append(w_element("jc", val="center"))
        p_pr.append(w_element("spacing", before="0", after="0", line="240", lineRule="auto"))
        p_pr.append(w_element("keepNext"))
        tc.append(paragraph)
    else:
        p = w_element("p")
        p_pr = w_element("pPr")
        p_pr.append(w_element("jc", val="right" if number_text else "center"))
        p_pr.append(w_element("spacing", before="0", after="0"))
        p.append(p_pr)
        if number_text is not None:
            run = w_element("r")
            run_pr = w_element("rPr")
            r_fonts = w_element("rFonts", ascii="Times New Roman", hAnsi="Times New Roman", eastAsia="宋体", cs="Times New Roman")
            run_pr.append(r_fonts)
            run_pr.append(w_element("sz", val="21"))
            run_pr.append(w_element("szCs", val="21"))
            run.append(run_pr)
            text = w_element("t")
            text.text = number_text
            run.append(text)
            p.append(run)
        tc.append(p)
    return tc


def make_equation_table(equation_paragraph, tag: str):
    table = w_element("tbl")
    table_pr = w_element("tblPr")
    table_pr.append(w_element("tblW", w="5000", type="pct"))
    table_pr.append(w_element("jc", val="center"))
    table_pr.append(w_element("tblLayout", type="fixed"))
    borders = w_element("tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(w_element(side, val="nil"))
    table_pr.append(borders)
    table.append(table_pr)

    grid = w_element("tblGrid")
    for width in (650, 8100, 650):
        grid.append(w_element("gridCol", w=str(width)))
    table.append(grid)

    row = w_element("tr")
    row_pr = w_element("trPr")
    row_pr.append(w_element("cantSplit"))
    row.append(row_pr)
    row.append(make_cell(650))
    row.append(make_cell(8100, paragraph=equation_paragraph))
    row.append(make_cell(650, number_text=f"({tag})"))
    table.append(row)
    return table


def rewrite_equation_numbers(path: Path, expected: int) -> int:
    with zipfile.ZipFile(path, "r") as source_zip:
        members = {name: source_zip.read(name) for name in source_zip.namelist()}

    root = etree.fromstring(members["word/document.xml"])
    body = root.find("w:body", namespaces=NS)
    if body is None:
        raise RuntimeError("word/document.xml has no w:body")

    markers: list[tuple[etree._Element, str]] = []
    for paragraph in root.xpath(".//w:p", namespaces=NS):
        text = "".join(paragraph.xpath(".//w:t/text()", namespaces=NS)).strip()
        if text.startswith(EQUATION_MARKER_PREFIX) and text.endswith(EQUATION_MARKER_SUFFIX):
            tag = text[len(EQUATION_MARKER_PREFIX) : -len(EQUATION_MARKER_SUFFIX)]
            tag = re.sub(r"^\d{4}_", "", tag)
            markers.append((paragraph, tag))

    if len(markers) != expected:
        raise RuntimeError(f"expected {expected} equation markers, found {len(markers)}")

    converted = 0
    for marker, tag in markers:
        equation_paragraph = marker.getprevious()
        while equation_paragraph is not None and equation_paragraph.tag != f"{{{W_NS}}}p":
            equation_paragraph = equation_paragraph.getprevious()
        if equation_paragraph is None or not equation_paragraph.xpath(".//m:oMathPara", namespaces=NS):
            raise RuntimeError(f"equation paragraph missing before marker {tag}")
        table = make_equation_table(equation_paragraph, tag)
        marker.getparent().replace(marker, table)
        converted += 1

    members["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    temp_path = path.with_suffix(".equations.tmp.docx")
    with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as destination_zip:
        for name, data in members.items():
            destination_zip.writestr(name, data)
    temp_path.replace(path)
    return converted


def inspect_docx(path: Path, stats: dict[str, int]) -> dict[str, int | bool | str]:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        xml = archive.read("word/document.xml")
        footer_xml = b"\n".join(
            archive.read(name)
            for name in names
            if name.startswith("word/footer") and name.endswith(".xml")
        )
    root = etree.fromstring(xml)
    visible_text = "\n".join(
        "".join(paragraph.xpath(".//w:t/text()", namespaces=NS)).strip()
        for paragraph in root.xpath(".//w:p", namespaces=NS)
    )
    visible_text = "\n".join(line for line in visible_text.splitlines() if line)
    equation_numbers = re.findall(r"^\((\d+(?:[a-z])?)\)$", visible_text, flags=re.MULTILINE)
    return {
        "docx_bytes": path.stat().st_size,
        "visible_text_chars": len(visible_text),
        "nonempty_paragraphs": sum(1 for line in visible_text.splitlines() if line.strip()),
        "native_math_objects": len(root.xpath(".//m:oMath", namespaces=NS)),
        "numbered_equation_tables": len(equation_numbers),
        "all_equation_tags_preserved": len(equation_numbers) == stats["equation_count"],
        "word_tables": len(root.xpath(".//w:tbl", namespaces=NS)),
        "embedded_media": sum(1 for name in names if name.startswith("word/media/")),
        "page_number_field": " PAGE " in footer_xml.decode("utf-8", errors="ignore"),
        "equation_marker_leftovers": visible_text.count(EQUATION_MARKER_PREFIX),
        "citation_processor_leftovers": len(re.findall(r"\[@C\d+", visible_text)),
        "raw_display_math_delimiters": visible_text.count("$$"),
        "raw_tag_commands": visible_text.count("\\tag{"),
        "working_citation_groups_normalized": stats["working_citation_groups"],
    }


def write_report(path: Path, source: Path, output: Path, stats: dict[str, int], audit: dict[str, int | bool | str]) -> None:
    required_pass = (
        bool(audit["all_equation_tags_preserved"])
        and int(audit["embedded_media"]) >= stats["image_count"]
        and int(audit["equation_marker_leftovers"]) == 0
        and int(audit["citation_processor_leftovers"]) == 0
        and int(audit["raw_display_math_delimiters"]) == 0
        and int(audit["raw_tag_commands"]) == 0
        and bool(audit["page_number_field"])
    )
    lines = [
        "# Word Export Report",
        "",
        f"- Status: {'PASS' if required_pass else 'FAIL'}",
        f"- Date: {date.today().isoformat()}",
        f"- Source: `{source}`",
        f"- Output: `{output}`",
        "- Layout: A4, 2.54 cm margins, single-column review manuscript",
        "- Body: Times New Roman/宋体 12 pt, 1.5 line spacing, justified",
        "- Headings: Times New Roman/黑体 with hierarchical sizing",
        "- Equations: native Word OMML, centred, right-aligned equation numbers",
        "- Tables: repeated shaded header rows, fixed widths, rows kept together",
        "- Figures: embedded and constrained to the printable page width",
        "- Citations: working keys retained as `[Cxxx]`; they are not final author-year references",
        "",
        "## Source structure",
        "",
    ]
    for key, value in stats.items():
        lines.append(f"- {key}: {value}")
    lines.extend(("", "## DOCX audit", ""))
    for key, value in audit.items():
        lines.append(f"- {key}: {value}")
    lines.extend(
        (
            "",
            "## Boundary",
            "",
            "- The DOCX is a formatted Chinese review manuscript, not the final publisher typeset proof.",
            "- Hidden HTML author comments are intentionally excluded from the visible Word manuscript.",
            "- The locked citation keys remain working identifiers pending the final verified bibliography replacement.",
            "",
        )
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines), encoding="utf-8")
    if not required_pass:
        raise RuntimeError("DOCX audit failed; see word_report.md")


def main() -> int:
    args = parse_args()
    source = args.source.resolve()
    output = args.output.resolve()
    report = args.report.resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    output.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="manuscript-zh-trc-docx-") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        processed_markdown = temp_dir / "manuscript_zh_trc_word_source.md"
        reference_docx = temp_dir / "reference.docx"
        intermediate_docx = temp_dir / "manuscript_zh_trc_intermediate.docx"

        stats = preprocess_markdown(source, processed_markdown)
        if stats["equation_count"] == 0:
            raise RuntimeError("no numbered display equations detected")
        build_reference_docx(reference_docx)

        reader = "markdown+tex_math_dollars+pipe_tables+implicit_figures+raw_tex-citations"
        pypandoc.convert_file(
            str(processed_markdown),
            "docx",
            format=reader,
            outputfile=str(intermediate_docx),
            extra_args=(
                f"--reference-doc={reference_docx}",
                f"--resource-path={OUTPUT_ROOT}",
                "--standalone",
                "--wrap=none",
                "--dpi=300",
                "--metadata=lang:zh-CN",
            ),
        )

        shutil.copy2(intermediate_docx, output)
        style_document(output)
        converted = rewrite_equation_numbers(output, stats["equation_count"])
        if converted != stats["equation_count"]:
            raise RuntimeError(f"converted {converted} equations, expected {stats['equation_count']}")

        audit = inspect_docx(output, stats)
        write_report(report, source, output, stats, audit)

    print(f"wrote {output}")
    print(f"wrote {report}")
    print(f"equations={stats['equation_count']}")
    print(f"tables={stats['markdown_table_count']}")
    print(f"images={stats['image_count']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
